from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


PROJECT_ROOT = Path(__file__).resolve().parent.parent
OUTPUT = PROJECT_ROOT / "reports" / "综合化工桌面管理系统实训报告.docx"


COLORS = {
    "heading": "2E74B5",
    "heading_dark": "1F4D78",
    "ink": "0B2545",
    "muted": "666666",
    "table_header": "F2F4F7",
    "callout": "F4F6F9",
    "border": "A6A6A6",
}


def set_font(run, size=None, bold=None, color=None, east_asia="Microsoft YaHei", ascii_font="Calibri"):
    run.font.name = ascii_font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: list[int]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    tbl_grid = tbl.tblGrid
    if tbl_grid is None:
        tbl_grid = OxmlElement("w:tblGrid")
        tbl.insert(0, tbl_grid)
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        tbl_grid.append(grid_col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.width = Inches(widths_dxa[index] / 1440)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[index]))
            tc_w.set(qn("w:type"), "dxa")


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name, size, color, before, after in [
        ("Title", 22, COLORS["ink"], 0, 12),
        ("Subtitle", 12, COLORS["muted"], 0, 14),
        ("Heading 1", 16, COLORS["heading"], 16, 8),
        ("Heading 2", 13, COLORS["heading"], 12, 6),
        ("Heading 3", 12, COLORS["heading_dark"], 8, 4),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    header_p = section.header.paragraphs[0]
    header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header_p.add_run("综合化工桌面管理系统实训报告")
    set_font(run, size=9, color=COLORS["muted"])

    footer_p = section.footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_p.add_run("题目10 · 项目模块化设计 · 权限逻辑 · 功能整合 · 全局异常处理")
    set_font(run, size=9, color=COLORS["muted"])


def add_title_block(doc: Document) -> None:
    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("综合化工桌面管理系统实训报告")
    set_font(run, size=22, bold=True, color=COLORS["ink"])

    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("题目10：整合数据查询、参数监控、工单管理、报表导出的小型桌面软件")
    set_font(run, size=12, color=COLORS["muted"])

    table = doc.add_table(rows=5, cols=2)
    table.style = "Table Grid"
    rows = [
        ("学生姓名", "连代平"),
        ("项目名称", "综合化工桌面管理系统"),
        ("开发语言", "Python 3 + Tkinter + SQLite"),
        ("工程位置", r"C:\Users\连代平\PycharmProjects\ChemicalManagementSystem"),
        ("完成日期", str(date.today())),
    ]
    for row, (label, value) in zip(table.rows, rows):
        row.cells[0].text = label
        row.cells[1].text = value
        set_cell_shading(row.cells[0], COLORS["table_header"])
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    set_font(run, size=10.5)
    set_table_geometry(table, [1900, 7460])

    doc.add_paragraph()


def add_paragraph(doc, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.10
    run = p.add_run(text)
    set_font(run)
    return p


def add_bullets(doc, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.15)
        run = p.add_run(item)
        set_font(run)


def add_numbered(doc, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(item)
        set_font(run)


def add_callout(doc, title: str, body: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    set_cell_shading(cell, COLORS["callout"])
    set_cell_margins(cell, top=140, bottom=140, start=180, end=180)
    p = cell.paragraphs[0]
    run = p.add_run(title + "：")
    set_font(run, bold=True, color=COLORS["heading_dark"])
    run = p.add_run(body)
    set_font(run)
    set_table_geometry(table, [9360])


def add_table(doc, headers: list[str], rows: list[list[str]], widths: list[int]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        set_cell_shading(cell, COLORS["table_header"])
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                set_font(run, bold=True, color=COLORS["ink"])
    for row_data in rows:
        row = table.add_row()
        for i, value in enumerate(row_data):
            cell = row.cells[i]
            cell.text = value
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                if len(value) < 8:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    set_font(run, size=10)
    set_table_geometry(table, widths)
    doc.add_paragraph()


def build_report() -> None:
    doc = Document()
    style_document(doc)
    add_title_block(doc)

    doc.add_heading("一、实习目标与任务概述", level=1)
    add_paragraph(
        doc,
        "本项目围绕题目10“综合化工桌面管理系统”展开，目标是综合运用项目模块化设计、权限逻辑、功能整合、工程规范和全局异常处理等知识点，完成一个完整可运行的小型桌面软件。系统面向化工企业日常管理场景，整合化学品台账、关键工艺参数、处理工单和报表输出等功能。",
    )
    add_callout(
        doc,
        "完成情况",
        "系统已实现菜单栏与工具栏布局、运行看板、数据查询、参数监控、工单管理、报表预览/导出、权限管理、操作审计和统一异常处理，并通过编译、单元测试与界面 smoke 验证。",
    )

    doc.add_heading("二、需求分析", level=1)
    add_table(
        doc,
        ["需求项", "实现说明", "完成状态"],
        [
            ["菜单栏+工具栏", "主窗口提供系统、模块、管理、帮助菜单，并提供常用模块工具栏入口。", "已完成"],
            ["数据查询", "按关键字、类别、危险等级查询化学品台账，管理员可入库/出库。", "已完成"],
            ["参数监控", "支持自动采样、趋势图、报警弹窗、报警历史、报警确认和阈值配置。", "已完成"],
            ["工单管理", "支持新建、编辑、状态流转、删除，普通操作员不能删除或关闭工单。", "已完成"],
            ["报表导出", "支持综合简报、库存、参数、工单 CSV 报表导出。", "已完成"],
            ["身份区分", "基于 SQLite 的用户、角色、权限表实现 RBAC 权限控制。", "已完成"],
            ["异常处理", "统一捕获用户操作异常，弹窗提示并写入日志文件。", "已完成"],
        ],
        [1800, 6200, 1360],
    )

    doc.add_heading("三、总体设计", level=1)
    add_paragraph(
        doc,
        "系统采用分层模块化结构。界面层使用 Tkinter 和 ttk 组件构建桌面窗口；业务层通过 DataStore、ReportService、Auth 等模块封装数据访问、报表和权限逻辑；数据层使用 SQLite 保存用户、角色、业务数据、历史数据和审计记录。",
    )
    add_table(
        doc,
        ["层次", "主要文件", "职责"],
        [
            ["程序入口", "main.py", "创建应用对象并启动主循环。"],
            ["界面层", "app/ui/*.py", "主窗口、看板、查询、监控、工单、报表、权限管理等视图。"],
            ["业务层", "app/data_store.py, app/auth.py, app/report_service.py", "业务数据读写、权限判断、报表预览和导出。"],
            ["数据层", "app/database.py, data/cms_data.db", "SQLite 建表、初始化和持久化存储。"],
            ["配置与异常", "app/config.py, app/exceptions.py", "路径配置、全局日志和异常弹窗处理。"],
            ["测试层", "tests/test_core.py", "覆盖认证、权限、库存、参数、工单、报表和审计逻辑。"],
        ],
        [1500, 3100, 4760],
    )

    doc.add_heading("四、数据库与权限设计", level=1)
    add_paragraph(
        doc,
        "早期版本使用 JSON 文件保存业务数据，后续升级为 SQLite 数据库存储。数据库首次启动时自动创建，默认数据由 seed_data.py 提供，保证工程在无外部依赖的情况下可直接运行。",
    )
    add_table(
        doc,
        ["数据表", "用途"],
        [
            ["users", "保存用户账号、密码、显示名称、角色和启用状态。"],
            ["roles", "保存角色编码、角色名称和说明。"],
            ["permissions", "保存系统权限点，例如库存调整、报警确认、角色管理。"],
            ["role_permissions", "维护角色与权限点之间的多对多关系。"],
            ["chemicals", "保存化学品台账、库存、库位、CAS、供应商等信息。"],
            ["inventory_history", "记录库存调整历史。"],
            ["parameters", "保存监控参数、阈值、当前值和确认状态。"],
            ["parameter_samples", "保存参数采样历史，用于趋势图展示。"],
            ["alarm_history", "保存报警发生、确认人和确认时间。"],
            ["work_orders", "保存工单台账和处理状态。"],
            ["audit_logs", "保存关键操作审计记录。"],
        ],
        [2500, 6860],
    )
    add_paragraph(
        doc,
        "权限控制采用 RBAC 思路：用户绑定角色，角色绑定权限点。系统启动和操作执行时均通过数据库权限表判断是否允许操作。管理员拥有完整权限，普通操作员默认只能查询、查看监控、新建/编辑工单和预览报表。",
    )

    doc.add_heading("五、功能模块实现", level=1)
    modules = [
        ("运行看板", "展示化学品数量、高风险物料、报警数量、未关闭工单；管理员可查看操作审计。"),
        ("数据查询", "支持台账筛选和危险等级标记；管理员可执行入库和出库，系统写入库存历史与审计日志。"),
        ("参数监控", "支持手动采样、自动刷新、实时趋势图、报警弹窗、报警历史、阈值配置和报警确认。"),
        ("工单管理", "支持工单新建、编辑、状态流转；管理员可删除和关闭工单，普通操作员受限。"),
        ("报表预览/导出", "支持综合简报、库存、参数、工单台账的预览和 CSV 导出。"),
        ("权限管理", "管理员可新增用户、重置密码、启停用户、分配角色、新增角色和勾选角色权限。"),
    ]
    for name, desc in modules:
        doc.add_heading(name, level=2)
        add_paragraph(doc, desc)

    doc.add_heading("六、关键技术实现", level=1)
    add_bullets(
        doc,
        [
            "使用 Tkinter/ttk 作为桌面 UI 框架，采用菜单栏、工具栏和多视图切换结构。",
            "使用 SQLite 标准库实现持久化数据存储，无需安装数据库服务。",
            "使用 RBAC 权限表实现用户、角色、权限配置，避免权限逻辑完全写死。",
            "使用统一 DataStore 接口封装数据库操作，降低界面层与数据层耦合。",
            "使用统一异常处理机制捕获界面回调异常，提升用户体验和系统稳定性。",
            "使用 CSV 导出报表，便于用 Excel 打开和提交。",
            "使用 unittest 覆盖核心逻辑，保证功能迭代后仍可验证。",
        ],
    )

    doc.add_heading("七、测试与验证", level=1)
    add_table(
        doc,
        ["测试内容", "验证点", "结果"],
        [
            ["语法编译", "python -m compileall app main.py tests", "通过"],
            ["单元测试", "python -m unittest discover -s tests，共 7 项", "通过"],
            ["SQLite 集成", "验证数据库初始化、权限读取、参数采样数据", "通过"],
            ["权限验证", "管理员与普通操作员的菜单、按钮和后台权限不同", "通过"],
            ["监控验证", "参数采样、趋势图数据、报警历史生成", "通过"],
            ["报表验证", "综合简报 CSV 可导出并包含表头", "通过"],
        ],
        [1900, 5700, 1760],
    )

    doc.add_heading("八、运行说明", level=1)
    add_numbered(
        doc,
        [
            r"打开项目目录：C:\Users\连代平\PycharmProjects\ChemicalManagementSystem。",
            "运行命令：python main.py，或在 PyCharm 中直接运行 main.py。",
            "管理员账号：admin，密码：admin123。",
            "普通操作员账号：operator，密码：operator123。",
            "首次启动会自动生成 SQLite 数据库 data/cms_data.db。",
            "导出的 CSV 报表位于 reports 目录，运行日志位于 logs/app.log。",
        ],
    )

    doc.add_heading("九、问题与改进", level=1)
    add_paragraph(
        doc,
        "本次开发过程中，最初版本使用 JSON 文件保存数据，权限也主要通过代码常量控制。后续根据完整工程要求，升级为 SQLite 数据库，并增加用户、角色、权限配置表，使系统更接近工业项目架构。参数监控也由单纯模拟采样升级为自动刷新、趋势图和报警历史。",
    )
    add_bullets(
        doc,
        [
            "后续可继续增加密码加密存储，提高账号安全性。",
            "可增加更完整的 Excel/PDF 报表模板，增强报表展示效果。",
            "可引入真实设备通讯协议或接口，使参数监控接入实际传感器数据。",
            "可增加安装包打包脚本，便于在其他电脑部署运行。",
        ],
    )

    doc.add_heading("十、实习总结", level=1)
    add_paragraph(
        doc,
        "通过本项目，完成了一个具备完整工程结构的小型桌面管理系统。系统不仅实现了题目要求的四大业务模块，还补充了运行看板、SQLite 数据库、角色权限配置、操作审计、参数趋势和报警历史等增强功能。开发过程体现了模块化设计、数据持久化、权限控制、异常处理和测试验证等软件工程思想，为后续开发更复杂的工业管理系统打下了基础。",
    )

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_report()
